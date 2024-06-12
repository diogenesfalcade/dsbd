from docx import Document
from deep_translator import GoogleTranslator
from tqdm import tqdm
from nltk.tokenize import sent_tokenize
import sys

inputFile = sys.argv[1]
outputFile =  inputFile.replace('input', 'output')

def translate_text(paragraph):
    """Wrapper for Google Translate with upload workaround.
    
    Collects chuncks of senteces below limit to translate.
    """
    # Set-up and wrap translation client
    translate = GoogleTranslator(source='en', target='pt').translate

    # Split input text into a list of sentences
    sentences = sent_tokenize(paragraph)

    # Initialize containers
    translated_text = ''
    source_text_chunk = ''

    # collect chuncks of sentences below limit and translate them individually
    for sentence in sentences:
        # if chunck together with current sentence is below limit, add the sentence
        if ((len(sentence.encode('utf-8')) + len(source_text_chunk.encode('utf-8')) < 5000)):
            source_text_chunk += ' ' + sentence
        
        # else translate chunck and start new one with current sentence
        else:
            translated_text += ' ' + translate(source_text_chunk)

            # if current sentence smaller than 5000 chars, start new chunck
            if (len(sentence.encode('utf-8')) < 5000):
                source_text_chunk = sentence

            # else, replace sentence with notification message
            else:    
                message = "<<Omitted Word longer than 5000bytes>>"
                translated_text += ' ' + translate(message)

                # Re-set text container to empty
                source_text_chunk = ''

    # Translate the final chunk of input text, if there is any valid text left to translate
    if translate(source_text_chunk) != None:
        translated_text += ' ' + translate(source_text_chunk)
    
    return translated_text

def process_docx(input_file, output_file):
    # Abrir o documento DOCX
    doc = Document(input_file)

    # Coletar todos os elementos de texto (parágrafos e células de tabela)
    elements = []
    for paragraph in doc.paragraphs:
        elements.append(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                elements.append(cell)
                
    # Iterar sobre os elementos e traduzir o texto com barra de progresso
    for index, element in enumerate(tqdm(elements, desc="Traduzindo documento", unit="element")):
        if element.text.strip() and element != '' and element.text:  # Verifica se o elemento não está vazio 
            try:
                translated_text = translate_text(element.text)
                element.text = translated_text
            except:
                element.text = "<<EXCEÇÃO>>"
                doc.save("output/erro.docx")


    # Iterar sobre os rodapés e traduzir o texto
    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            if paragraph.text.strip():  # Verifica se o parágrafo não está vazio
                translated_text = translate_text(paragraph.text)
                paragraph.text = translated_text

    # Salvar o documento traduzido
    doc.save(output_file)

if __name__ == "__main__":
    process_docx(inputFile, outputFile)
    print(f"Tradução concluída. Arquivo salvo como {outputFile}")
