from docx import Document
from deep_translator import GoogleTranslator
from tqdm import tqdm
from nltk.tokenize import sent_tokenize
from multiprocessing import Pool
import sys

inputFile = sys.argv[1]
outputFile = inputFile.replace('input', 'output')

def translate_text(paragraph):
    """Wrapper for Google Translate with upload workaround.
    
    Collects chunks of sentences below limit to translate.
    """
    # Set-up and wrap translation client
    translate = GoogleTranslator(source='en', target='pt').translate

    # Split input text into a list of sentences
    sentences = sent_tokenize(paragraph)

    # Initialize containers
    translated_text = ''
    source_text_chunk = ''

    # Collect chunks of sentences below limit and translate them individually
    for sentence in sentences:
        # If chunk together with current sentence is below limit, add the sentence
        if (len(sentence.encode('utf-8')) + len(source_text_chunk.encode('utf-8')) < 5000):
            source_text_chunk += ' ' + sentence
        else:
            translated_text += ' ' + translate(source_text_chunk)

            # If current sentence is smaller than 5000 chars, start new chunk
            if (len(sentence.encode('utf-8')) < 5000):
                source_text_chunk = sentence
            else:
                message = "<<Omitted Word longer than 5000bytes>>"
                translated_text += ' ' + translate(message)
                source_text_chunk = ''

    # Translate the final chunk of input text, if there is any valid text left to translate
    if source_text_chunk:
        translated_text += ' ' + translate(source_text_chunk)

    return translated_text

def process_text(text):
    """Translate a single piece of text."""
    if text.strip():
        try:
            translated_text = translate_text(text)
            return translated_text
        except Exception as e:
            print(f"Error translating text: {e}")
            return "<<EXCEÇÃO>>"
    return text

def contains_image(paragraph):
    """Check if a paragraph contains an image."""
    for run in paragraph.runs:
        if 'graphic' in run._element.xml:
            return True
    return False

def process_docx(input_file, output_file):
    # Open the document
    doc = Document(input_file)

    # Collect all the text elements (paragraphs and table cells)
    elements = []
    for paragraph in doc.paragraphs:
        elements.append(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                elements.append(cell)

    # Extract texts from elements, checking for images
    texts = [element.text if not isinstance(element, Document) and not contains_image(element) else None for element in elements]

    # Translate texts in parallel, skipping elements with images
    with Pool() as pool:
        translated_texts = list(tqdm(pool.imap(process_text, (text for text in texts if text is not None)), total=sum(1 for text in texts if text is not None), desc="Traduzindo documento", unit="element"))

    translated_texts_iter = iter(translated_texts)

    # Update elements with translated texts, skipping elements with images
    for element, original_text in zip(elements, texts):
        if original_text is not None:
            element.text = next(translated_texts_iter)

    # Iterate over footnotes
    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            if paragraph.text.strip():  # Check if paragraph is not empty
                translated_text = translate_text(paragraph.text)
                paragraph.text = translated_text

    # Save file
    doc.save(output_file)

if __name__ == "__main__":
    process_docx(inputFile, outputFile)
    print(f"Tradução concluída. Arquivo salvo como {outputFile}")
